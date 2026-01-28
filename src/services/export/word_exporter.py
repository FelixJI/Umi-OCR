#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Umi-OCR Word导出器

将OCR结果导出为Word格式。

Author: Umi-OCR Team
Date: 2026-01-27
"""

import logging
from typing import List, Dict, Any
from docx import Document

from .base_exporter import BaseExporter

logger = logging.getLogger(__name__)


class WordExporter(BaseExporter):
    """Word导出器"""

    def __init__(self):
        """初始化Word导出器"""
        logger.info("Word导出器初始化完成")

    def export(self, data: List[Dict[str, Any]], output_path: str, **kwargs) -> bool:
        """
        导出为 Word

        Args:
            data: OCR 结果列表
            output_path: 输出文件路径
            **kwargs: 额外参数

        Returns:
            bool: 是否成功
        """
        try:
            doc = Document()

            for item in data:
                text = item.get("text", "")
                title = item.get("title", "")

                if title:
                    doc.add_heading(title, level=1)

                doc.add_paragraph(text)
                doc.add_page_break()

            doc.save(output_path)
            logger.info(f"Word 导出成功: {output_path}")

            return True

        except Exception as e:
            logger.error(f"Word导出失败: {e}", exc_info=True)
            return False

    def get_file_extension(self) -> str:
        """获取文件扩展名"""
        return ".docx"

    def get_name(self) -> str:
        """获取导出器名称"""
        return "Word"
